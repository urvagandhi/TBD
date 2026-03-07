"""
DOCX Writer — Consumes docx_instructions from TRANSFORM agent.

Six dedicated builders + in-place mode:
  1. build_apa_docx():       APA-specific — page-based sections
  2. build_ieee_docx():      IEEE-specific — 2-column, 10pt, single spacing
  3. build_springer_docx():  Springer Nature — 10pt, single spacing, justified, numeric headings
  4. build_chicago_docx():   Chicago 17th Ed — 12pt, double spacing, left-aligned, un-numbered headings
  5. build_vancouver_docx(): Vancouver/ICMJE — 12pt, double spacing, UPPERCASE H1, numbered citations
  6. write_formatted_docx(): Generic fallback — rules-driven flat sections
  7. transform_docx_in_place(): In-place transformation of uploaded DOCX files.

Routing is handled by crew.py via style_key from detect_style().
"""
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from tools.logger import get_logger
from tools.tool_errors import DocumentWriteError

logger = get_logger(__name__)

_CASE_OPTIONS = {"Title Case", "UPPERCASE", "Sentence case", "lowercase"}


def _media_lookup(store: Optional[dict], key) -> Optional[dict]:
    """Look up media by key, trying both original type and int conversion."""
    if not store or key == "" or key is None:
        return None
    result = store.get(key)
    if result:
        return result
    try:
        return store.get(int(key))
    except (ValueError, TypeError):
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# APA-SPECIFIC DOCX BUILDER (from APA Pipeline Prompts §6)
# ═══════════════════════════════════════════════════════════════════════════════

def build_apa_docx(transform_output: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """Main entry point: takes TRANSFORM agent JSON → produces APA DOCX file."""

    instructions = transform_output.get("docx_instructions", transform_output)
    sections_data = instructions.get("sections", [])

    if not sections_data:
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    doc = Document()

    # ── 1. SET DOCUMENT-LEVEL DEFAULTS ──
    style = doc.styles['Normal']
    font = style.font
    font.name = instructions.get("font", "Times New Roman")
    # Support both field names: font_size_halfpoints (MD prompt) and font_size (legacy)
    if "font_size_halfpoints" in instructions:
        font.size = Pt(instructions["font_size_halfpoints"] / 2)
    else:
        font.size = Pt(instructions.get("font_size", 12))

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 2.0
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set default font in XML
    rpr = doc.styles['Normal'].element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rpr.append(rFonts)

    # ── 2. CONFIGURE HEADING STYLES ──
    _configure_heading_styles(doc)

    # ── 3. PROCESS EACH SECTION ──
    first_section = True
    for section_data in sections_data:
        section_type = section_data.get("type", "body")

        if section_type == "title_page":
            _write_title_page(doc, section_data, instructions, first_section)
        elif section_type == "abstract_page":
            _write_abstract_page(doc, section_data, instructions)
        elif section_type == "body":
            _write_body(doc, section_data, instructions, image_store=image_store, table_store=table_store)
        elif section_type == "references_page":
            _write_references_page(doc, section_data, instructions)
        else:
            _write_body(doc, section_data, instructions, image_store=image_store, table_store=table_store)

        first_section = False

    # ── 4. SET PAGE SIZE & MARGINS ON ALL SECTIONS ──
    page_size = instructions.get("page_size", {})
    margins = instructions.get("margins", {})
    for section in doc.sections:
        section.page_width = Twips(page_size.get("width", 12240))
        section.page_height = Twips(page_size.get("height", 15840))
        section.top_margin = Twips(margins.get("top", 1440))
        section.bottom_margin = Twips(margins.get("bottom", 1440))
        section.left_margin = Twips(margins.get("left", 1440))
        section.right_margin = Twips(margins.get("right", 1440))

        _add_page_number_header(section)

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_APA] Written | file=%s | sections=%d", out.name, len(sections_data))
    return str(out)


def _configure_heading_styles(doc):
    """Set up APA heading styles in the document."""
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing = 2.0
    h1.paragraph_format.first_line_indent = Inches(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.first_line_indent = Inches(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 2.0
    h3.paragraph_format.first_line_indent = Inches(0.5)

    # H4: Bold, 0.5" indent, Title Case, inline with text, ends with period
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Times New Roman'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = False
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.space_before = Pt(0)
    h4.paragraph_format.space_after = Pt(0)
    h4.paragraph_format.line_spacing = 2.0
    h4.paragraph_format.first_line_indent = Inches(0.5)

    # H5: Bold Italic, 0.5" indent, Title Case, inline with text, ends with period
    h5 = doc.styles['Heading 5']
    h5.font.name = 'Times New Roman'
    h5.font.size = Pt(12)
    h5.font.bold = True
    h5.font.italic = True
    h5.font.color.rgb = RGBColor(0, 0, 0)
    h5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h5.paragraph_format.space_before = Pt(0)
    h5.paragraph_format.space_after = Pt(0)
    h5.paragraph_format.line_spacing = 2.0
    h5.paragraph_format.first_line_indent = Inches(0.5)


def _write_title_page(doc, section_data, instructions, is_first):
    """Write the APA title page."""
    if not is_first:
        doc.add_section()

    # Default 3 blank lines if no explicit spacing element comes first
    elements = section_data.get("elements", [])
    has_leading_spacing = elements and elements[0].get("type") == "spacing"
    if not has_leading_spacing:
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    for element in elements:
        etype = element.get("type", "")

        if etype == "spacing":
            # Blank lines from the MD prompt schema
            blank_lines = element.get("blank_lines", 1)
            for _ in range(blank_lines):
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

        elif etype == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(element.get("text", ""))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype in ("authors", "affiliation"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(element.get("text", ""))
            run.bold = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def _write_abstract_page(doc, section_data, instructions):
    """Write the APA abstract page."""
    doc.add_section()

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "abstract_label":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", "Abstract"))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "abstract_body":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            _add_text_with_italics(p, element.get("text", ""))

        elif etype == "keywords":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            # Support both first_line_indent: true/720 and explicit values
            fi = element.get("first_line_indent", True)
            if fi is True or (isinstance(fi, (int, float)) and fi > 0):
                p.paragraph_format.first_line_indent = Inches(0.5)
            else:
                p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("label", "Keywords: "))
            run.italic = element.get("label_italic", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Support both field names: "items" (MD prompt) and "keywords" (legacy)
            keywords = element.get("items", element.get("keywords", []))
            if isinstance(keywords, list):
                keywords_text = ", ".join(str(k) for k in keywords)
            else:
                keywords_text = str(keywords)
            run2 = p.add_run(keywords_text)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)


def _write_body(doc, section_data, instructions, image_store=None, table_store=None):
    """Write all body content (intro through discussion)."""
    doc.add_section()
    # Support both field names: body_first_line_indent_dxa (MD prompt) and body_first_line_indent (legacy)
    indent_dxa = instructions.get("body_first_line_indent_dxa", instructions.get("body_first_line_indent", 720))
    indent = Inches(indent_dxa / 1440) if indent_dxa > 13 else Inches(indent_dxa)

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "title_repeat":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", ""))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "heading":
            level = element.get("level", 1)
            level = min(max(level, 1), 5)
            is_inline = element.get("inline_with_text", False) or level >= 3
            following_text = element.get("following_text", "")

            if is_inline and level >= 3:
                # APA H3-H5: inline heading — heading and body text in SAME paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Inches(0.5)

                heading_text = element.get("text", "")
                if not heading_text.endswith("."):
                    heading_text += "."

                h_run = p.add_run(heading_text)
                h_run.bold = True
                h_run.font.name = 'Times New Roman'
                h_run.font.size = Pt(12)
                h_run.font.color.rgb = RGBColor(0, 0, 0)
                if level in (3, 5):
                    h_run.italic = True

                if following_text:
                    space_run = p.add_run("  ")
                    space_run.font.name = 'Times New Roman'
                    space_run.font.size = Pt(12)
                    _add_text_with_italics(p, following_text)
            else:
                # APA H1/H2: separate paragraph headings
                capped_level = min(level, 3)
                p = doc.add_heading(element.get("text", ""), level=capped_level)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        elif etype == "body_paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = indent
            _add_text_with_italics(p, element.get("text", ""))

        elif etype in ("figure_caption", "figure_block"):
            _render_figure_block_apa(doc, element, image_store)

        elif etype in ("table_caption", "table_block"):
            _render_table_block_apa(doc, element, table_store)


def _render_figure_block_apa(doc, element: dict, image_store: Optional[dict] = None) -> None:
    """
    APA §7.4 compliant figure block:
      [Image — centered]
      Figure N         ← bold, flush left
      Caption text     ← italic, flush left
    """
    fig_num = element.get("number", "")
    image_data = _media_lookup(image_store, fig_num)

    # Step 1: Insert actual image if available
    if image_data and image_data.get("bytes"):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.line_spacing = 2.0
        img_para.paragraph_format.space_before = Pt(12)
        img_para.paragraph_format.first_line_indent = Inches(0)
        run = img_para.add_run()

        # Scale to max text width (6.5" for US Letter with 1" margins)
        max_width = 6.5
        img_w = image_data.get("width", 0)
        img_h = image_data.get("height", 0)

        if img_w > 0:
            # For DOCX sources, dimensions are in EMU (914400 EMU = 1 inch)
            if "width_emu" in image_data:
                display_width = min(max_width, image_data["width_emu"] / 914400)
            else:
                display_width = min(max_width, img_w / 96)
        else:
            display_width = max_width

        try:
            run.add_picture(BytesIO(image_data["bytes"]), width=Inches(display_width))
            logger.info("[DOCX] Inserted Figure %s image (%dx%d)", fig_num, img_w, img_h)
        except Exception as e:
            logger.warning("[DOCX] Failed to insert Figure %s image: %s", fig_num, e)
            run.text = f"[Figure {fig_num} — image could not be rendered]"
            run.italic = True
    else:
        # Placeholder when no image was extracted
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder.paragraph_format.line_spacing = 2.0
        placeholder.paragraph_format.first_line_indent = Inches(0)
        run = placeholder.add_run(f"[Figure {fig_num} — image not available in source]")
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Step 2: "Figure N" — bold, flush left
    label_para = doc.add_paragraph()
    label_para.paragraph_format.line_spacing = 2.0
    label_para.paragraph_format.first_line_indent = Inches(0)
    label = element.get("label", f"Figure {fig_num}")
    label_run = label_para.add_run(label)
    label_run.bold = True
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(12)

    # Step 3: Caption — italic, flush left
    caption = element.get("caption", "")
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.paragraph_format.line_spacing = 2.0
        cap_para.paragraph_format.first_line_indent = Inches(0)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = "Times New Roman"
        cap_run.font.size = Pt(12)

    # Blank line after figure block
    doc.add_paragraph()


def _render_table_block_apa(doc, element: dict, table_store: Optional[dict] = None) -> None:
    """
    APA §7.22 compliant table block:
      Table N         ← bold, flush left
      Caption text    ← italic, flush left  (ABOVE table)
      [Table data]
      Note.           ← if present
    """
    tbl_num = element.get("number", "")
    table_data = _media_lookup(table_store, tbl_num)

    # Step 1: "Table N" — bold, flush left (ABOVE table per APA)
    label_para = doc.add_paragraph()
    label_para.paragraph_format.line_spacing = 2.0
    label_para.paragraph_format.first_line_indent = Inches(0)
    label = element.get("label", f"Table {tbl_num}")
    label_run = label_para.add_run(label)
    label_run.bold = True
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(12)

    # Step 2: Caption — italic, flush left (also ABOVE table)
    caption = element.get("caption", "")
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.paragraph_format.line_spacing = 2.0
        cap_para.paragraph_format.first_line_indent = Inches(0)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = "Times New Roman"
        cap_run.font.size = Pt(12)

    # Step 3: Render actual table if data available
    rows = None
    if table_data and table_data.get("rows"):
        rows = table_data["rows"]
    elif element.get("rows"):
        rows = element["rows"]

    if rows and len(rows) > 0:
        try:
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = "Table Grid"

            for r_idx, row_data in enumerate(rows):
                row = table.rows[r_idx]
                for c_idx in range(num_cols):
                    cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                    cell = row.cells[c_idx]
                    cell.text = str(cell_text) if cell_text else ""
                    # Bold header row
                    if r_idx == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)
                    else:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)

            logger.info("[DOCX] Inserted Table %s with %d rows × %d cols", tbl_num, len(rows), num_cols)
        except Exception as e:
            logger.warning("[DOCX] Failed to render Table %s: %s", tbl_num, e)
            p = doc.add_paragraph(f"[Table {tbl_num} — could not render table data]")
            p.runs[0].italic = True
    else:
        p = doc.add_paragraph(f"[Table {tbl_num} — data not available in source]")
        p.paragraph_format.first_line_indent = Inches(0)
        p.runs[0].italic = True
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(12)

    # Blank line after table block
    doc.add_paragraph()


def _write_references_page(doc, section_data, instructions):
    """Write the references page with hanging indent."""
    doc.add_section()

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "references_label":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", "References"))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "reference_entry":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            _add_text_with_italics(p, element.get("text", ""))


def _add_text_with_italics(paragraph, text, font_name='Times New Roman', font_size_pt=12, base_bold=False):
    """Parse *italic*, **bold**, and ***bold+italic*** markers and create appropriate runs."""
    if not text:
        return
    # Match ***bold+italic***, **bold**, or *italic* markers
    parts = re.split(r'(\*{1,3}[^*]+\*{1,3})', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if base_bold:
                run.bold = True
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)


def _add_page_number_header(section):
    """Add right-aligned page number to section header."""
    header = section.header
    header.is_linked_to_previous = False

    if not header.paragraphs:
        p = header.add_paragraph()
    else:
        p = header.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════════════════════
# IEEE-SPECIFIC DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_ieee_docx(instructions: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """
    IEEE-specific DOCX builder.

    Handles flat section lists (title, heading, paragraph, reference, etc.)
    with IEEE rules: 2-column layout, 10pt Times New Roman, single spacing,
    numbered references, specific margin/heading conventions from rules/ieee.json.
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])
    rules = instructions.get("rules", {}) or {}

    # Template-aware document generation
    style_name = (rules.get("style_name") or "IEEE").strip()
    template_path = Path(__file__).parent.parent / "templates" / f"{style_name}.docx"

    if template_path.exists():
        logger.info("[DOCX_IEEE] Using journal template: %s", template_path.name)
        doc = Document(str(template_path))
    else:
        logger.debug("[DOCX_IEEE] No template found for '%s' — using blank document", style_name)
        doc = Document()

    doc_rules = rules.get("document", {})
    font_name = doc_rules.get("font", "Times New Roman")
    font_size = _safe_int(doc_rules.get("font_size", 10), 10)
    line_spacing = _safe_float(doc_rules.get("line_spacing", 1.0), 1.0)
    margins = doc_rules.get("margins", {})
    columns = _safe_int(doc_rules.get("columns", 2), 2)
    doc_alignment = doc_rules.get("alignment", "justify")

    _apply_document_defaults(doc, font_name, font_size, line_spacing)
    _set_document_margins(doc, margins)
    
    # Force first section to 1 column for Title/Authors
    _set_columns(doc.sections[0], 1)
    
    in_two_column = False

    fig_rules = rules.get("figures", {})
    tbl_rules = rules.get("tables", {})

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            is_front_matter = section_type in ("title", "authors", "affiliations", "author_blocks")
            if not is_front_matter and not in_two_column and columns > 1:
                # Transition to 2 columns
                new_sect = doc.add_section(WD_SECTION.CONTINUOUS)
                _set_columns(new_sect, columns)
                in_two_column = True

            if section_type == "title":
                title_rules = rules.get("title_page", {})
                _add_title(doc, content, title_rules, font_name, font_size)
            elif section_type == "authors":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                # Split multiple authors by newline if provided, or comma
                lines = content.split('\n')
                for i, line in enumerate(lines):
                    run = para.add_run(line.strip())
                    _apply_font(run, font_name, font_size)
                    if i < len(lines) - 1:
                        para.add_run('\n')
            elif section_type == "author_blocks":
                blocks = section.get("blocks", [])
                if blocks:
                    _add_author_blocks_table(doc, blocks, font_name, font_size, line_spacing)
            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                heading_rules = rules.get("headings", {}).get(f"H{level}", {})
                _add_heading(doc, content, level, heading_rules, font_name, font_size)
            elif section_type == "abstract":
                abstract_rules = rules.get("abstract", {})
                _add_abstract_inline(doc, content, abstract_rules, font_name, font_size, line_spacing, default_align=doc_alignment)
            elif section_type == "keywords":
                kw_rules = rules.get("keywords_section", rules.get("abstract", {}))
                _add_keywords(doc, content, kw_rules, font_name, font_size)
            elif section_type == "reference":
                ref_rules = rules.get("references", {})
                _add_reference(doc, content, ref_rules, font_name, font_size)
            elif section_type in ("figure_caption", "figure_block"):
                _render_figure_image(doc, section, image_store, font_name, font_size)
                _add_ieee_figure_caption(doc, section, fig_rules, font_name, font_size, line_spacing)
            elif section_type in ("table_caption", "table_block"):
                _add_ieee_table_caption(doc, section, tbl_rules, font_name, font_size, line_spacing)
                _render_table_data(doc, section, table_store, font_name, font_size)
            else:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, doc_alignment)
        except Exception as e:
            logger.warning("[DOCX_IEEE] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, doc_alignment)
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_IEEE] Written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


def _render_figure_image(
    doc: Document, section: dict, image_store: Optional[dict],
    font_name: str, font_size: int,
) -> None:
    """Insert actual image above caption (shared by IEEE/generic writers)."""
    fig_num = section.get("number", "")
    image_data = _media_lookup(image_store, fig_num)
    if not image_data or not image_data.get("bytes"):
        return

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(6)
    run = img_para.add_run()

    max_width = 3.0  # Narrower for 2-column layouts
    img_w = image_data.get("width", 0)
    if "width_emu" in image_data:
        display_width = min(max_width, image_data["width_emu"] / 914400)
    elif img_w > 0:
        display_width = min(max_width, img_w / 96)
    else:
        display_width = max_width

    try:
        run.add_picture(BytesIO(image_data["bytes"]), width=Inches(display_width))
        logger.info("[DOCX] Inserted Figure %s image", fig_num)
    except Exception as e:
        logger.warning("[DOCX] Failed to insert Figure %s image: %s", fig_num, e)


def _render_table_data(
    doc: Document, section: dict, table_store: Optional[dict],
    font_name: str, font_size: int,
) -> None:
    """Insert actual table data below caption (shared by IEEE/generic writers)."""
    tbl_num = section.get("number", "")
    table_data = _media_lookup(table_store, tbl_num)

    rows = None
    if table_data and table_data.get("rows"):
        rows = table_data["rows"]
    elif section.get("rows"):
        rows = section["rows"]

    if not rows:
        return

    try:
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(num_cols):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                cell = row.cells[c_idx]
                cell.text = str(cell_text) if cell_text else ""
                for para in cell.paragraphs:
                    for run in para.runs:
                        _apply_font(run, font_name, font_size, bold=(r_idx == 0))

        logger.info("[DOCX] Inserted Table %s with %d rows × %d cols", tbl_num, len(rows), num_cols)
    except Exception as e:
        logger.warning("[DOCX] Failed to render Table %s: %s", tbl_num, e)


def _add_ieee_figure_caption(
    doc: Document, section: dict, fig_rules: dict,
    font_name: str, font_size: int, line_spacing: float,
) -> None:
    """
    IEEE figure caption: "Fig. 1." label + caption text.

    IEEE conventions (from rules/ieee.json):
      - label_prefix: "Fig."  (not "Figure")
      - label NOT bold, NOT italic
      - caption NOT italic, centered, below figure
      - numbering: arabic (1, 2, 3...)
    """
    label_prefix = fig_rules.get("label_prefix", "Fig.")
    label_bold = fig_rules.get("label_bold", False)
    label_italic = fig_rules.get("label_italic", False)
    caption_italic = fig_rules.get("caption_italic", False)
    caption_align = fig_rules.get("caption_alignment", "center")

    number = section.get("number", "")
    caption = section.get("caption", section.get("content", ""))
    label_text = f"{label_prefix} {number}." if number else f"{label_prefix}"

    para = doc.add_paragraph()
    if caption_align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif caption_align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _apply_line_spacing(para.paragraph_format, line_spacing)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Label run: "Fig. 1."
    label_run = para.add_run(label_text + " ")
    _apply_font(label_run, font_name, font_size, bold=label_bold, italic=label_italic)

    # Caption run
    if caption:
        cap_run = para.add_run(caption)
        _apply_font(cap_run, font_name, font_size, bold=False, italic=caption_italic)


def _add_ieee_table_caption(
    doc: Document, section: dict, tbl_rules: dict,
    font_name: str, font_size: int, line_spacing: float,
) -> None:
    """
    IEEE table caption: "TABLE I" label (above table) + caption below.

    IEEE conventions (from rules/ieee.json):
      - label_prefix: "TABLE"  (UPPERCASE)
      - label bold, NOT italic
      - caption centered, above table
      - numbering: roman (I, II, III...)
      - border_style: full_grid
    """
    label_prefix = tbl_rules.get("label_prefix", "TABLE")
    label_bold = tbl_rules.get("label_bold", True)
    label_italic = tbl_rules.get("label_italic", False)
    caption_italic = tbl_rules.get("caption_italic", False)
    caption_align = tbl_rules.get("caption_alignment", "center")
    numbering = tbl_rules.get("numbering", "roman")

    number = section.get("number", "")
    caption = section.get("caption", section.get("content", ""))

    # Convert to roman numeral if needed
    if numbering == "roman" and number:
        number = _to_roman(number)

    label_text = f"{label_prefix} {number}" if number else label_prefix

    # Line 1: "TABLE I" (label)
    label_para = doc.add_paragraph()
    if caption_align == "center":
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_line_spacing(label_para.paragraph_format, line_spacing)
    label_para.paragraph_format.space_before = Pt(6)
    label_para.paragraph_format.space_after = Pt(0)
    label_run = label_para.add_run(label_text)
    _apply_font(label_run, font_name, font_size, bold=label_bold, italic=label_italic)

    # Line 2: caption text
    if caption:
        cap_para = doc.add_paragraph()
        if caption_align == "center":
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_line_spacing(cap_para.paragraph_format, line_spacing)
        cap_para.paragraph_format.space_before = Pt(0)
        cap_para.paragraph_format.space_after = Pt(6)
        cap_run = cap_para.add_run(caption)
        _apply_font(cap_run, font_name, font_size, bold=False, italic=caption_italic)


def _to_roman(number) -> str:
    """Convert integer or string number to roman numeral."""
    try:
        n = int(number)
    except (TypeError, ValueError):
        return str(number)
    vals = [(1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')]
    result = ""
    for val, sym in vals:
        while n >= val:
            result += sym
            n -= val
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# SPRINGER NATURE WRITER (sn-mathphys-ay, author-date)
# ═══════════════════════════════════════════════════════════════════════════════

def build_springer_docx(instructions: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """
    Springer Nature DOCX builder.

    Handles flat section lists with Springer rules:
    - 10pt Times New Roman, single spacing, justified
    - Margins: 1" top/bottom, 1.25" left/right
    - Numeric heading hierarchy (1, 1.1, 1.1.1)
    - Author-date citations: (Smith 2020)
    - Alphabetical references with hanging indent
    - Fig. N (bold) captions below, Table N (bold) captions above
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])
    rules = instructions.get("rules", {}) or {}

    doc = Document()

    # Springer: 10pt Times New Roman, single spacing, justified
    font_name = "Times New Roman"
    font_size = 10
    line_spacing = 1.0

    _apply_document_defaults(doc, font_name, font_size, line_spacing)

    # Margins: 1" top/bottom, 1.25" left/right
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    # Configure Springer heading styles
    _configure_springer_heading_styles(doc, font_name)

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            if section_type == "title":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, 14, bold=True)

            elif section_type == "authors":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            elif section_type == "affiliations":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, 9, italic=True)

            elif section_type == "abstract":
                # Label
                label_para = doc.add_paragraph()
                label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(label_para.paragraph_format, line_spacing)
                label_run = label_para.add_run("Abstract")
                _apply_font(label_run, font_name, font_size, bold=True)
                # Body
                body_para = doc.add_paragraph()
                body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _apply_line_spacing(body_para.paragraph_format, line_spacing)
                _add_text_with_italics(body_para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type == "keywords":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run("Keywords ")
                _apply_font(label_run, font_name, font_size, bold=True)
                text_run = para.add_run(content)
                _apply_font(text_run, font_name, font_size)

            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                bold = section.get("bold", level <= 2)
                italic = section.get("italic", level >= 3)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size, bold=bold, italic=italic)

            elif section_type == "reference":
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type in ("figure_caption", "figure_block"):
                _render_figure_image(doc, section, image_store, font_name, font_size)
                # Springer: "Fig. N" bold, caption below
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run(f"Fig. {number}. " if number else "Fig. ")
                _apply_font(label_run, font_name, font_size, bold=True)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)

            elif section_type in ("table_caption", "table_block"):
                # Springer: "Table N" bold, caption above
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run(f"Table {number} " if number else "Table ")
                _apply_font(label_run, font_name, font_size, bold=True)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)
                _render_table_data(doc, section, table_store, font_name, font_size)

            else:
                # Body paragraph: justified, 10pt, single spacing
                _add_paragraph(doc, content, font_name, font_size, line_spacing, "justify")

        except Exception as e:
            logger.warning("[DOCX_SPRINGER] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, "justify")
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_SPRINGER] Written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


def _configure_springer_heading_styles(doc, font_name="Times New Roman"):
    """Set up Springer heading styles: H1/H2 bold, H3 italic, all numbered."""
    h1 = doc.styles['Heading 1']
    h1.font.name = font_name
    h1.font.size = Pt(10)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    h2 = doc.styles['Heading 2']
    h2.font.name = font_name
    h2.font.size = Pt(10)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    h3 = doc.styles['Heading 3']
    h3.font.name = font_name
    h3.font.size = Pt(10)
    h3.font.bold = False
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


# ═══════════════════════════════════════════════════════════════════════════════
# CHICAGO 17TH EDITION WRITER (Author-Date)
# ═══════════════════════════════════════════════════════════════════════════════

def build_chicago_docx(instructions: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """
    Chicago Manual of Style (17th Ed, Author-Date) DOCX builder.

    Handles flat section lists with Chicago rules:
    - 12pt Times New Roman, double spacing, left-aligned (ragged right)
    - 1" margins all sides
    - First-line indent 0.5" for body paragraphs
    - Un-numbered headings: H1 centered/bold, H2 left/not bold, H3 left/italic
    - Author-date citations: (Smith 2020)
    - Alphabetical references with hanging indent 0.5"
    - Figure N. caption below, Table N caption above
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])

    doc = Document()

    # Chicago: 12pt Times New Roman, double spacing, left-aligned
    font_name = "Times New Roman"
    font_size = 12
    line_spacing = 2.0

    _apply_document_defaults(doc, font_name, font_size, line_spacing)

    # 1" margins all sides
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Configure Chicago heading styles
    _configure_chicago_heading_styles(doc, font_name)

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            if section_type == "title":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)  # Not bold per Chicago

            elif section_type == "authors":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            elif section_type == "affiliations":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            elif section_type == "abstract":
                # Chicago: "Abstract" centered, not bold
                label_para = doc.add_paragraph()
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(label_para.paragraph_format, line_spacing)
                label_run = label_para.add_run("Abstract")
                _apply_font(label_run, font_name, font_size)
                # Body: left-aligned, first-line indent
                body_para = doc.add_paragraph()
                body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                body_para.paragraph_format.first_line_indent = Inches(0.5)
                _apply_line_spacing(body_para.paragraph_format, line_spacing)
                _add_text_with_italics(body_para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type == "keywords":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run("Keywords: ")
                _apply_font(label_run, font_name, font_size, italic=True)
                text_run = para.add_run(content)
                _apply_font(text_run, font_name, font_size)

            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                para = doc.add_paragraph()
                _apply_line_spacing(para.paragraph_format, line_spacing)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(0)
                if level == 1:
                    # H1: Centered, Bold, Title Case
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(content)
                    _apply_font(run, font_name, font_size, bold=True)
                elif level == 2:
                    # H2: Left-aligned, Title Case, Not bold
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = para.add_run(content)
                    _apply_font(run, font_name, font_size)
                else:
                    # H3+: Left-aligned, Italic, Title Case
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = para.add_run(content)
                    _apply_font(run, font_name, font_size, italic=True)

            elif section_type == "reference":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type in ("figure_caption", "figure_block"):
                _render_figure_image(doc, section, image_store, font_name, font_size)
                # Chicago: "Figure N. Caption" below, left-aligned
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_text = f"Figure {number}. " if number else "Figure. "
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)

            elif section_type in ("table_caption", "table_block"):
                # Chicago: "Table N" + caption above, minimal borders
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_text = f"Table {number}" if number else "Table"
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size)
                if caption:
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _apply_line_spacing(cap_para.paragraph_format, line_spacing)
                    cap_run = cap_para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)
                _render_table_data(doc, section, table_store, font_name, font_size)

            else:
                # Body paragraph: left-aligned, first-line indent 0.5"
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0.5)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

        except Exception as e:
            logger.warning("[DOCX_CHICAGO] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, "left")
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_CHICAGO] Written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


def _configure_chicago_heading_styles(doc, font_name="Times New Roman"):
    """Set up Chicago heading styles: H1 centered/bold, H2 left, H3 left/italic."""
    h1 = doc.styles['Heading 1']
    h1.font.name = font_name
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing = 2.0

    h2 = doc.styles['Heading 2']
    h2.font.name = font_name
    h2.font.size = Pt(12)
    h2.font.bold = False
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing = 2.0

    h3 = doc.styles['Heading 3']
    h3.font.name = font_name
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 2.0


# ═══════════════════════════════════════════════════════════════════════════════
# VANCOUVER (ICMJE / BIOMEDICAL) WRITER
# ═══════════════════════════════════════════════════════════════════════════════

def build_vancouver_docx(instructions: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """
    Vancouver (ICMJE) DOCX builder for biomedical manuscripts.

    Handles flat section lists with Vancouver rules:
    - 12pt Times New Roman, double spacing, left-aligned
    - 1" margins all sides
    - First-line indent 0.5" for body paragraphs
    - Structured headings: H1 UPPERCASE/bold, H2 Title Case/bold, H3 italic
    - Numbered citations [1], ordered by first appearance
    - References numbered by appearance, hanging indent
    - Figure N. caption below, Table N. caption above
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])

    doc = Document()

    # Vancouver: 12pt Times New Roman, double spacing, left-aligned
    font_name = "Times New Roman"
    font_size = 12
    line_spacing = 2.0

    _apply_document_defaults(doc, font_name, font_size, line_spacing)

    # 1" margins all sides
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Configure Vancouver heading styles
    _configure_vancouver_heading_styles(doc, font_name)

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            if section_type == "title":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)  # Not bold per Vancouver

            elif section_type == "authors":
                # Vancouver: Surname Initials (no periods)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            elif section_type == "affiliations":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            elif section_type == "abstract":
                # Vancouver: "Abstract" bold, left-aligned
                label_para = doc.add_paragraph()
                label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(label_para.paragraph_format, line_spacing)
                label_run = label_para.add_run("Abstract")
                _apply_font(label_run, font_name, font_size, bold=True)
                # Body: left-aligned, no indent for abstract
                body_para = doc.add_paragraph()
                body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(body_para.paragraph_format, line_spacing)
                _add_text_with_italics(body_para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type == "keywords":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run("Keywords: ")
                _apply_font(label_run, font_name, font_size, bold=True)
                text_run = para.add_run(content)
                _apply_font(text_run, font_name, font_size)

            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(0)
                if level == 1:
                    # H1: Bold, UPPERCASE
                    run = para.add_run(content.upper())
                    _apply_font(run, font_name, font_size, bold=True)
                elif level == 2:
                    # H2: Bold, Title Case
                    run = para.add_run(content)
                    _apply_font(run, font_name, font_size, bold=True)
                else:
                    # H3+: Italic, Title Case
                    run = para.add_run(content)
                    _apply_font(run, font_name, font_size, italic=True)

            elif section_type == "reference":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

            elif section_type in ("figure_caption", "figure_block"):
                _render_figure_image(doc, section, image_store, font_name, font_size)
                # Vancouver: "Figure N. Caption" below figure
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_text = f"Figure {number}. " if number else "Figure. "
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size, bold=True)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)

            elif section_type in ("table_caption", "table_block"):
                # Vancouver: "Table N. Caption" above table
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_text = f"Table {number}. " if number else "Table. "
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size, bold=True)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size)
                _render_table_data(doc, section, table_store, font_name, font_size)

            else:
                # Body paragraph: left-aligned, first-line indent 0.5"
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0.5)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

        except Exception as e:
            logger.warning("[DOCX_VANCOUVER] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, "left")
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_VANCOUVER] Written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


def _configure_vancouver_heading_styles(doc, font_name="Times New Roman"):
    """Set up Vancouver heading styles: H1 bold/uppercase, H2 bold, H3 italic."""
    h1 = doc.styles['Heading 1']
    h1.font.name = font_name
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing = 2.0

    h2 = doc.styles['Heading 2']
    h2.font.name = font_name
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing = 2.0

    h3 = doc.styles['Heading 3']
    h3.font.name = font_name
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 2.0


# ═══════════════════════════════════════════════════════════════════════════════
# GENERIC WRITER (fallback for other formats)
# ═══════════════════════════════════════════════════════════════════════════════

def write_formatted_docx(instructions: dict, output_path: str, image_store: Optional[dict] = None, table_store: Optional[dict] = None) -> str:
    """
    Fully rules-driven generic DOCX writer for custom / unknown journals.

    Every formatting decision is read from the rules JSON (passed via
    instructions["rules"]) so that when a user provides a custom PDF and
    custom rules are generated, those rules drive the output exactly.

    Rules keys consumed:
      document      → font, font_size, line_spacing, margins, columns, alignment
      title_page    → title_bold, title_centered, title_font_size, title_case
      abstract      → label, label_bold, label_centered, label_italic,
                       indent_first_line, keywords_label, keywords_italic
      headings.H1/2/3 → bold, italic, centered, case, font_size, underline
      references    → hanging_indent, indent_size, line_spacing, label_bold,
                       label_centered, space_between_entries
      figures       → label_prefix, label_bold, caption_italic, caption_alignment
      tables        → label_prefix, label_bold, caption_italic, caption_alignment
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])

    # ── All formatting comes from rules ──
    rules = instructions.get("rules", {}) or {}

    # Template-aware document generation
    style_name = (rules.get("style_name") or "Standard").strip()
    template_path = Path(__file__).parent.parent / "templates" / f"{style_name}.docx"

    if template_path.exists():
        logger.info("[DOCX] Using journal template: %s", template_path.name)
        doc = Document(str(template_path))
    else:
        logger.debug("[DOCX] No template found for '%s' — using blank document", style_name)
        doc = Document()

    # ── Document-level rules ──
    doc_rules = rules.get("document", {})
    font_name = doc_rules.get("font", "Times New Roman")
    font_size = _safe_int(doc_rules.get("font_size", 12), 12)
    line_spacing = _safe_float(doc_rules.get("line_spacing", 2.0), 2.0)
    margins = doc_rules.get("margins", {})
    columns = _safe_int(doc_rules.get("columns", 1), 1)
    doc_alignment = doc_rules.get("alignment", "justify")

    _apply_document_defaults(doc, font_name, font_size, line_spacing)
    _set_document_margins(doc, margins)
    _set_columns(doc, columns)

    # ── Pre-fetch sub-rule dicts once ──
    title_rules = rules.get("title_page", {})
    abstract_rules = rules.get("abstract", {})
    heading_rules_map = rules.get("headings", {})
    ref_rules = rules.get("references", {})
    fig_rules = rules.get("figures", {})
    tbl_rules = rules.get("tables", {})

    # Reference-specific line spacing (may differ from body)
    ref_line_spacing = _safe_float(ref_rules.get("line_spacing", line_spacing), line_spacing)
    ref_space_between = ref_rules.get("space_between_entries", False)

    # Alignment helper
    def _resolve_alignment(align_str: str):
        a = align_str.lower()
        if a == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if a == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        if a == "justify":
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return WD_ALIGN_PARAGRAPH.LEFT

    body_align = _resolve_alignment(doc_alignment)

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            # ── Title ──
            if section_type == "title":
                para = doc.add_paragraph()
                centered = title_rules.get("title_centered", True)
                bold = title_rules.get("title_bold", False)
                title_size = _safe_int(title_rules.get("title_font_size", font_size), font_size)
                title_case = title_rules.get("title_case", "")
                display = _apply_case_transform(content, title_case) if title_case else content
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(display)
                _apply_font(run, font_name, title_size, bold=bold)

            # ── Authors ──
            elif section_type == "authors":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if title_rules.get("title_centered", True) else WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            # ── Affiliations ──
            elif section_type == "affiliations":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if title_rules.get("title_centered", True) else WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                run = para.add_run(content)
                _apply_font(run, font_name, font_size)

            # ── Abstract ──
            elif section_type == "abstract":
                label = abstract_rules.get("label", "Abstract")
                label_bold = abstract_rules.get("label_bold", False)
                label_italic = abstract_rules.get("label_italic", False)
                label_centered = abstract_rules.get("label_centered", True)
                indent_first = abstract_rules.get("indent_first_line", False)

                label_para = doc.add_paragraph()
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if label_centered else WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(label_para.paragraph_format, line_spacing)
                label_run = label_para.add_run(label)
                _apply_font(label_run, font_name, font_size, bold=label_bold, italic=label_italic)

                body_para = doc.add_paragraph()
                body_para.alignment = body_align
                if indent_first:
                    body_para.paragraph_format.first_line_indent = Inches(0.5)
                _apply_line_spacing(body_para.paragraph_format, line_spacing)
                _add_text_with_italics(body_para, content, font_name=font_name, font_size_pt=font_size)

            # ── Keywords ──
            elif section_type == "keywords":
                kw_label = abstract_rules.get("keywords_label", "Keywords:")
                kw_italic = abstract_rules.get("keywords_italic", False)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run(kw_label + " ")
                _apply_font(label_run, font_name, font_size, italic=kw_italic)
                text_run = para.add_run(content)
                _apply_font(text_run, font_name, font_size)

            # ── Headings (fully rules-driven per level) ──
            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                h_rules = heading_rules_map.get(f"H{level}", {})
                bold = h_rules.get("bold", True)
                italic = h_rules.get("italic", False)
                centered = h_rules.get("centered", False)
                underline = h_rules.get("underline", False)
                case = h_rules.get("case", "Title Case")
                h_font_size = _safe_int(h_rules.get("font_size", font_size), font_size)
                display = _apply_case_transform(content, case) if case else content
                # Apply UPPERCASE from case rule
                if case == "UPPERCASE":
                    display = display.upper()

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
                _apply_line_spacing(para.paragraph_format, line_spacing)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(display)
                _apply_font(run, font_name, h_font_size, bold=bold, italic=italic)
                if underline:
                    run.underline = True

            # ── References (rules-driven spacing + indent) ──
            elif section_type == "reference":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hanging = ref_rules.get("hanging_indent", True)
                indent_size = _parse_measurement(ref_rules.get("indent_size", 0.5))
                if hanging:
                    para.paragraph_format.left_indent = Inches(indent_size)
                    para.paragraph_format.first_line_indent = Inches(-indent_size)
                _apply_line_spacing(para.paragraph_format, ref_line_spacing)
                if ref_space_between:
                    para.paragraph_format.space_after = Pt(6)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

            # ── Figures (rules-driven label + position) ──
            elif section_type in ("figure_caption", "figure_block"):
                label_prefix = fig_rules.get("label_prefix", "Figure")
                label_bold = fig_rules.get("label_bold", True)
                caption_italic = fig_rules.get("caption_italic", False)
                cap_align = fig_rules.get("caption_alignment", "center")
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                label_text = f"{label_prefix} {number}. " if number else f"{label_prefix} "

                _render_figure_image(doc, section, image_store, font_name, font_size)
                para = doc.add_paragraph()
                para.alignment = _resolve_alignment(cap_align)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size, bold=label_bold)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size, italic=caption_italic)

            # ── Tables (rules-driven label + position) ──
            elif section_type in ("table_caption", "table_block"):
                label_prefix = tbl_rules.get("label_prefix", "Table")
                label_bold = tbl_rules.get("label_bold", True)
                caption_italic = tbl_rules.get("caption_italic", False)
                cap_align = tbl_rules.get("caption_alignment", "center")
                number = section.get("number", "")
                caption = section.get("caption", section.get("content", ""))
                label_text = f"{label_prefix} {number}. " if number else f"{label_prefix} "

                para = doc.add_paragraph()
                para.alignment = _resolve_alignment(cap_align)
                _apply_line_spacing(para.paragraph_format, line_spacing)
                label_run = para.add_run(label_text)
                _apply_font(label_run, font_name, font_size, bold=label_bold)
                if caption:
                    cap_run = para.add_run(caption)
                    _apply_font(cap_run, font_name, font_size, italic=caption_italic)
                _render_table_data(doc, section, table_store, font_name, font_size)

            # ── Body paragraphs ──
            else:
                para = doc.add_paragraph()
                para.alignment = body_align
                _apply_line_spacing(para.paragraph_format, line_spacing)
                _add_text_with_italics(para, content, font_name=font_name, font_size_pt=font_size)

        except Exception as e:
            logger.warning("[DOCX_GENERIC] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing, doc_alignment)
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_GENERIC] Written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


# ═══════════════════════════════════════════════════════════════════════════════
# LEGACY HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_document_defaults(doc: Document, font_name: str, font_size: int, line_spacing: float) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        pf = style.paragraph_format
        _apply_line_spacing(pf, line_spacing)
    except Exception as e:
        logger.warning("[DOCX] Could not set document defaults: %s", e)


def _set_columns(doc_or_section, columns: int) -> None:
    """Set the number of columns for the document or section."""
    if columns <= 1:
        return
    try:
        if hasattr(doc_or_section, "sections"):
            section = doc_or_section.sections[0]
        else:
            section = doc_or_section
            
        cols_list = section._sectPr.xpath('./w:cols')
        if cols_list:
            cols = cols_list[0]
        else:
            cols = OxmlElement('w:cols')
            section._sectPr.append(cols)
        cols.set(qn('w:num'), str(columns))
    except Exception as e:
        logger.warning("[DOCX] Could not set columns: %s", e)


def _set_document_margins(doc: Document, margins: dict) -> None:
    for section in doc.sections:
        try:
            top = _parse_measurement(margins.get("top", 1.0))
            bottom = _parse_measurement(margins.get("bottom", 1.0))
            left = _parse_measurement(margins.get("left", 1.0))
            right = _parse_measurement(margins.get("right", 1.0))
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
        except Exception as e:
            logger.warning("[DOCX] Could not set margins: %s", e)


def _parse_measurement(value) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        value = value.strip().lower()
        try:
            if value.endswith("in"):
                return float(value[:-2])
            elif value.endswith("cm"):
                return float(value[:-2]) / 2.54
            elif value.endswith("mm"):
                return float(value[:-2]) / 25.4
            elif value.endswith("pt"):
                return float(value[:-2]) / 72.0
            else:
                return float(value)
        except ValueError:
            logger.warning("[DOCX] Cannot parse measurement '%s' — defaulting to 1.0in", value)
            return 1.0
    return 1.0


def _apply_line_spacing(pf, line_spacing: float) -> None:
    if line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        # Use MULTIPLE for all other values including 2.0
        # MULTIPLE scales with font size (correct for APA double-spacing)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing


def _apply_font(run, font_name: str, font_size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic


def _apply_case_transform(text: str, case_rule: str) -> str:
    if not case_rule or case_rule not in _CASE_OPTIONS:
        return text
    try:
        if case_rule == "UPPERCASE":
            return text.upper()
        elif case_rule == "lowercase":
            return text.lower()
        elif case_rule == "Sentence case":
            if not text:
                return text
            return text[0].upper() + text[1:].lower() if len(text) > 1 else text.upper()
        elif case_rule == "Title Case":
            _SMALL_WORDS = {
                "a", "an", "the", "and", "but", "or", "for", "nor",
                "on", "at", "to", "by", "in", "of", "up", "as", "is",
            }
            words = text.split()
            result = []
            for i, word in enumerate(words):
                if len(word) >= 2 and word.isupper():
                    result.append(word)
                elif "-" in word:
                    result.append("-".join(
                        part.capitalize() for part in word.split("-")
                    ))
                elif i == 0 or word.lower() not in _SMALL_WORDS:
                    result.append(word.capitalize())
                else:
                    result.append(word.lower())
            return " ".join(result)
    except Exception as e:
        logger.warning("[DOCX] Case transform failed for '%s': %s", case_rule, e)
    return text


def _add_paragraph(doc: Document, text: str, font_name: str, font_size: int, line_spacing: float, alignment: str = "justify") -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _apply_font(run, font_name, font_size)
    
    if alignment == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    _apply_line_spacing(para.paragraph_format, line_spacing)


def _add_title(doc: Document, text: str, title_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    centered = title_rules.get("title_centered", True)
    bold = title_rules.get("title_bold", False)
    size = title_rules.get("title_font_size", 24)

    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = para.add_run(text)
    _apply_font(run, font_name, size, bold=bold)


def _add_heading(doc: Document, text: str, level: int, heading_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = heading_rules.get("bold", True)
    italic = heading_rules.get("italic", False)
    centered = heading_rules.get("centered", False)
    case = heading_rules.get("case", "Title Case")
    heading_font_size = _safe_int(heading_rules.get("font_size", font_size), font_size)
    text = _apply_case_transform(text, case)
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _apply_font(run, font_name, heading_font_size, bold=bold, italic=italic)


def _add_abstract(doc: Document, text: str, abstract_rules: dict, font_name: str, font_size: int, line_spacing: float) -> None:
    label = abstract_rules.get("label", "Abstract")
    label_bold = abstract_rules.get("label_bold", False)
    label_italic = abstract_rules.get("label_italic", False)
    centered = abstract_rules.get("label_centered", True)

    label_para = doc.add_paragraph()
    if centered:
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_para.add_run(label)
    _apply_font(label_run, font_name, font_size, bold=label_bold, italic=label_italic)

    body_para = doc.add_paragraph()
    body_run = body_para.add_run(text)
    _apply_font(body_run, font_name, font_size)
    _apply_line_spacing(body_para.paragraph_format, line_spacing)


def _add_abstract_inline(doc: Document, text: str, abstract_rules: dict, font_name: str, font_size: int, line_spacing: float, default_align="justify") -> None:
    """Add abstract with an inline label (e.g. 'Abstract—Text...'), used for IEEE/Springer."""
    label = abstract_rules.get("label", "Abstract")
    if not label.endswith("—") and not label.endswith("-"):
        label += "—" # IEEE style usually has em-dash
    
    label_bold = abstract_rules.get("label_bold", True)
    label_italic = abstract_rules.get("label_italic", False)

    para = doc.add_paragraph()
    if default_align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif default_align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    _apply_line_spacing(para.paragraph_format, line_spacing)

    label_run = para.add_run(label)
    _apply_font(label_run, font_name, font_size, bold=label_bold, italic=label_italic)

    body_run = para.add_run(text)
    _apply_font(body_run, font_name, font_size)


def _add_keywords(doc: Document, text: str, rules: dict, font_name: str, font_size: int) -> None:
    label = rules.get("keywords_label", "Index Terms—")
    keywords_italic = rules.get("keywords_italic", False)
    para = doc.add_paragraph()
    label_run = para.add_run(label + " ")
    _apply_font(label_run, font_name, font_size, italic=keywords_italic)
    text_run = para.add_run(text)
    _apply_font(text_run, font_name, font_size)


def _add_reference(doc: Document, text: str, ref_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    hanging = ref_rules.get("hanging_indent", True)
    hanging_inches = _parse_measurement(ref_rules.get("indent_size", ref_rules.get("hanging_indent_inches", 0.5)))
    if hanging:
        para.paragraph_format.left_indent = Inches(hanging_inches)
        para.paragraph_format.first_line_indent = Inches(-hanging_inches)
    _add_text_with_italics(para, text, font_name=font_name, font_size_pt=font_size)


def _add_figure_caption(doc: Document, text: str, fig_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = fig_rules.get("label_bold", True)
    italic = fig_rules.get("caption_italic", False)
    alignment = fig_rules.get("caption_alignment", "center")

    if alignment == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(text)
    _apply_font(run, font_name, font_size, bold=bold, italic=italic)


def _add_table_caption(doc: Document, text: str, tbl_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = tbl_rules.get("label_bold", True)
    alignment = tbl_rules.get("caption_alignment", "center")

    if alignment == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(text)
    _apply_font(run, font_name, font_size, bold=bold)


def _add_author_blocks_table(doc: Document, blocks: list, font_name: str, font_size: int, line_spacing: float) -> None:
    """Add authors side-by-side using borderless tables. Wraps to max 3 columns."""
    if not blocks:
        return
        
    if len(blocks) == 4:
        chunks = [blocks[0:2], blocks[2:4]]
    else:
        max_cols = 3
        chunks = [blocks[i:i + max_cols] for i in range(0, len(blocks), max_cols)]
        
    for chunk_idx, chunk in enumerate(chunks):
        num_cols = len(chunk)
        table = doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = True
        
        for i, block in enumerate(chunk):
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_line_spacing(para.paragraph_format, line_spacing)
            
            name = block.get("name", "")
            if name:
                run = para.add_run(name)
                _apply_font(run, font_name, font_size, bold=False)

            affiliation = block.get("affiliation", [])
            if isinstance(affiliation, str):
                affiliation = [affiliation]
            
            for j, affil_line in enumerate(affiliation):
                if para.text:
                    para.add_run('\n')
                run = para.add_run(affil_line)
                # Typically the last line of affiliation (location) is regular, others italic
                is_italic = True
                if len(affiliation) > 1 and j == len(affiliation) - 1:
                    is_italic = False
                _apply_font(run, font_name, font_size, italic=is_italic)
                
            email = block.get("email", "")
            if email:
                if para.text:
                    para.add_run('\n')
                run = para.add_run(email)
                _apply_font(run, font_name, font_size, italic=False)
        
        # Add slight paragraph spacing between author rows if multiple rows exist
        if chunk_idx < len(chunks) - 1:
            spacer = doc.add_paragraph()
            _apply_line_spacing(spacer.paragraph_format, line_spacing)
            
    # Add a blank line after all authors
    para = doc.add_paragraph()
    _apply_line_spacing(para.paragraph_format, line_spacing)


# ═══════════════════════════════════════════════════════════════════════════════
# IN-PLACE DOCX TRANSFORMATION (for uploaded DOCX files)
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_heading_style(paragraph, fix: dict) -> None:
    bold = fix.get("bold", True)
    italic = fix.get("italic", False)
    centered = fix.get("centered", False)
    case = fix.get("case", "Title Case")
    text = paragraph.text
    text = _apply_case_transform(text, case)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
    if text != paragraph.text and paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""


def transform_docx_in_place(
    source_docx_path: str,
    transform_data: dict,
    rules: dict,
    output_path: str,
) -> str:
    """Apply journal formatting transformations to a source DOCX in-place."""
    try:
        doc = Document(source_docx_path)
    except Exception as e:
        raise DocumentWriteError(f"Cannot open source DOCX '{source_docx_path}': {e}") from e

    doc_rules      = rules.get("document", {})
    headings_rules = rules.get("headings", {})
    ref_rules      = rules.get("references", {})

    font_name       = doc_rules.get("font", "Times New Roman")
    font_size_pt    = _safe_int(doc_rules.get("font_size", 12), 12)
    line_spacing_v  = _safe_float(doc_rules.get("line_spacing", 2.0), 2.0)
    margins         = doc_rules.get("margins", {})

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(font_size_pt)
        _apply_line_spacing(normal.paragraph_format, line_spacing_v)
    except Exception as e:
        logger.warning("[DOCX_INPLACE] Could not set Normal style: %s", e)

    for sec in doc.sections:
        try:
            sec.top_margin    = Inches(_parse_measurement(margins.get("top",    1.0)))
            sec.bottom_margin = Inches(_parse_measurement(margins.get("bottom", 1.0)))
            sec.left_margin   = Inches(_parse_measurement(margins.get("left",   1.0)))
            sec.right_margin  = Inches(_parse_measurement(margins.get("right",  1.0)))
        except Exception as e:
            logger.warning("[DOCX_INPLACE] Could not set margins: %s", e)

    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            continue
        try:
            level = int(para.style.name.split()[-1])
        except (ValueError, IndexError):
            level = 1

        h_rules  = headings_rules.get(f"H{level}", {})
        case     = h_rules.get("case", None)
        bold     = h_rules.get("bold", True)
        italic   = h_rules.get("italic", False)
        centered = h_rules.get("centered", False)
        h_size   = _safe_int(h_rules.get("font_size", font_size_pt), font_size_pt)

        if centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if case:
            new_text = _apply_case_transform(para.text, case)
            if new_text != para.text:
                _replace_paragraph_text_preserve_runs(para, new_text)

        for run in para.runs:
            try:
                run.bold        = bold
                run.italic      = italic
                run.font.name   = font_name
                run.font.size   = Pt(h_size)
            except Exception as e:
                logger.warning("[DOCX_INPLACE] Heading run format error: %s", e)

    reference_order = transform_data.get("reference_order", [])
    _reorder_references_in_place(doc, reference_order, ref_rules)

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
    except Exception as e:
        raise DocumentWriteError(f"Failed to save in-place DOCX to '{out}': {e}") from e

    logger.info("[DOCX_INPLACE] Saved — file=%s", out.name)
    return str(out)


def _replace_paragraph_text_preserve_runs(para, new_text: str) -> None:
    if not para.runs:
        para.text = new_text
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _reorder_references_in_place(doc: Document, reference_order: list, ref_rules: dict) -> None:
    paragraphs = doc.paragraphs
    ref_start_idx: Optional[int] = None
    for i, para in enumerate(paragraphs):
        if re.search(r"^\s*references?\s*$", para.text, re.IGNORECASE):
            ref_start_idx = i
            break

    if ref_start_idx is None:
        return

    ref_paras = []
    for para in paragraphs[ref_start_idx + 1:]:
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip():
            ref_paras.append(para)

    if not ref_paras:
        return

    if reference_order and len(reference_order) == len(ref_paras):
        ordered = reference_order
    else:
        ordering = ref_rules.get("ordering", "alphabetical")
        if ordering == "alphabetical":
            ordered = sorted([p.text for p in ref_paras], key=str.lower)
        else:
            ordered = [p.text for p in ref_paras]

    for para, new_text in zip(ref_paras, ordered):
        if para.text != new_text:
            _replace_paragraph_text_preserve_runs(para, new_text)

    logger.info("[DOCX_INPLACE] References reordered — %d entries", len(ref_paras))


def _safe_int(value, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _safe_float(value, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default
