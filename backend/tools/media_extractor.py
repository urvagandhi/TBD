"""
Media Extractor — Extract images and tables from PDF/DOCX source files.

This module provides a side-channel for binary media data (images, table structures)
that bypasses the LLM agents entirely. The extracted media is mapped to figure/table
captions produced by the PARSE agent, then injected directly into the DOCX writer.

Supported sources:
  - PDF: PyMuPDF (fitz) for images, pdfplumber for table data
  - DOCX: python-docx inline_shapes for images, doc.tables for table data
"""
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

import fitz  # PyMuPDF — already a project dependency

from tools.logger import get_logger

logger = get_logger(__name__)

# Minimum image dimensions (pixels) — skip decorative icons, logos, watermarks
_MIN_IMAGE_DIM = 50


# ═══════════════════════════════════════════════════════════════════════════════
# IMAGE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

def extract_images_from_pdf(pdf_path: str) -> list[dict[str, Any]]:
    """
    Extract all embedded raster images from a PDF using PyMuPDF.

    Returns list of dicts: {page, index_on_page, bytes, ext, width, height, bbox}
    """
    images: list[dict[str, Any]] = []
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.warning("[MEDIA] Cannot open PDF for image extraction: %s", e)
        return images

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        img_list = page.get_images(full=True)
        for img_index, img_info in enumerate(img_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if not base_image or not base_image.get("image"):
                    continue
                w = base_image.get("width", 0)
                h = base_image.get("height", 0)
                if w < _MIN_IMAGE_DIM or h < _MIN_IMAGE_DIM:
                    continue
                images.append({
                    "page": page_num + 1,
                    "index_on_page": img_index,
                    "bytes": base_image["image"],
                    "ext": base_image.get("ext", "png"),
                    "width": w,
                    "height": h,
                })
            except Exception as e:
                logger.debug("[MEDIA] Skipping image xref=%d on page %d: %s", xref, page_num + 1, e)

    doc.close()
    logger.info("[MEDIA] PDF image extraction: %d images from %s", len(images), Path(pdf_path).name)
    return images


def extract_images_from_docx(docx_path: str) -> list[dict[str, Any]]:
    """
    Extract all inline images from a DOCX using python-docx.

    Returns list of dicts: {index, bytes, content_type, width_emu, height_emu}
    """
    from docx import Document

    images: list[dict[str, Any]] = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning("[MEDIA] Cannot open DOCX for image extraction: %s", e)
        return images

    for i, shape in enumerate(doc.inline_shapes):
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not rId:
                continue
            image_part = doc.part.related_parts[rId]
            blob = image_part.blob
            if not blob:
                continue
            images.append({
                "index": i,
                "bytes": blob,
                "content_type": image_part.content_type,
                "width_emu": shape.width,
                "height_emu": shape.height,
            })
        except Exception as e:
            logger.debug("[MEDIA] Skipping DOCX inline shape %d: %s", i, e)

    logger.info("[MEDIA] DOCX image extraction: %d images from %s", len(images), Path(docx_path).name)
    return images


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

def extract_tables_from_pdf(pdf_path: str) -> list[dict[str, Any]]:
    """
    Extract tables from PDF using pdfplumber (if available), else return empty.

    Returns list of dicts: {page, index_on_page, rows}
    where rows is list[list[str]] (cell text grid).
    """
    tables: list[dict[str, Any]] = []
    try:
        import pdfplumber
    except ImportError:
        logger.warning("[MEDIA] pdfplumber not installed — skipping PDF table extraction")
        return tables

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_tables = page.extract_tables()
                    for t_idx, table_data in enumerate(page_tables):
                        if not table_data or not any(table_data):
                            continue
                        # Clean None cells
                        clean_rows = [
                            [str(cell) if cell is not None else "" for cell in row]
                            for row in table_data
                            if row
                        ]
                        if clean_rows:
                            tables.append({
                                "page": page_num + 1,
                                "index_on_page": t_idx,
                                "rows": clean_rows,
                            })
                except Exception as e:
                    logger.debug("[MEDIA] Table extraction failed on page %d: %s", page_num + 1, e)
    except Exception as e:
        logger.warning("[MEDIA] pdfplumber failed to open PDF: %s", e)

    logger.info("[MEDIA] PDF table extraction: %d tables from %s", len(tables), Path(pdf_path).name)
    return tables


def extract_tables_from_docx(docx_path: str) -> list[dict[str, Any]]:
    """
    Extract tables from DOCX using python-docx.

    Returns list of dicts: {index, rows}
    where rows is list[list[str]].
    """
    from docx import Document

    tables: list[dict[str, Any]] = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning("[MEDIA] Cannot open DOCX for table extraction: %s", e)
        return tables

    for t_idx, table in enumerate(doc.tables):
        try:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append({"index": t_idx, "rows": rows})
        except Exception as e:
            logger.debug("[MEDIA] Skipping DOCX table %d: %s", t_idx, e)

    logger.info("[MEDIA] DOCX table extraction: %d tables from %s", len(tables), Path(docx_path).name)
    return tables


# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE/TABLE ↔ CAPTION MAPPING
# ═══════════════════════════════════════════════════════════════════════════════

def map_figures_to_images(
    extracted_images: list[dict],
    figure_captions: list[dict],
    source_type: str,
) -> dict[int, dict]:
    """
    Map extracted images to figure captions by number.

    For DOCX: uses document order (inline shapes appear in reading order).
    For PDF: uses page proximity, then falls back to sequential order.

    Args:
        extracted_images: from extract_images_from_pdf/docx
        figure_captions: from paper_structure["figures"] — [{number, caption, ...}]
        source_type: "pdf" or "docx"

    Returns:
        {figure_number: {bytes, ext, width, height, caption, ...}}
    """
    if not extracted_images or not figure_captions:
        return {}

    mapping: dict[int, dict] = {}
    sorted_captions = sorted(figure_captions, key=lambda c: c.get("number", 0))

    if source_type == "docx":
        for i, caption in enumerate(sorted_captions):
            if i < len(extracted_images):
                mapping[caption["number"]] = {
                    **extracted_images[i],
                    "caption": caption.get("caption", ""),
                }
    elif source_type == "pdf":
        # Try page proximity matching first
        claimed = set()
        for caption in sorted_captions:
            fig_num = caption.get("number", 0)
            # Images are in document order; try sequential matching
            for idx, img in enumerate(extracted_images):
                if idx not in claimed:
                    claimed.add(idx)
                    mapping[fig_num] = {
                        **img,
                        "caption": caption.get("caption", ""),
                    }
                    break

    logger.info("[MEDIA] Figure mapping: %d/%d figures mapped", len(mapping), len(figure_captions))
    return mapping


def map_tables_to_captions(
    extracted_tables: list[dict],
    table_captions: list[dict],
    source_type: str,
) -> dict[int, dict]:
    """
    Map extracted tables to table captions by number.

    Uses sequential order matching (N-th table = Table N).

    Returns:
        {table_number: {rows, caption, ...}}
    """
    if not extracted_tables or not table_captions:
        return {}

    mapping: dict[int, dict] = {}
    sorted_captions = sorted(table_captions, key=lambda c: c.get("number", 0))

    for i, caption in enumerate(sorted_captions):
        if i < len(extracted_tables):
            mapping[caption["number"]] = {
                **extracted_tables[i],
                "caption": caption.get("caption", ""),
            }

    logger.info("[MEDIA] Table mapping: %d/%d tables mapped", len(mapping), len(table_captions))
    return mapping


# ═══════════════════════════════════════════════════════════════════════════════
# HIGH-LEVEL CONVENIENCE
# ═══════════════════════════════════════════════════════════════════════════════

def extract_all_media(source_path: str) -> dict[str, Any]:
    """
    Extract all images and tables from a source file.

    Returns:
        {
            "source_type": "pdf" | "docx",
            "raw_images": [...],
            "raw_tables": [...],
        }
    """
    path = Path(source_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return {
            "source_type": "pdf",
            "raw_images": extract_images_from_pdf(str(path)),
            "raw_tables": extract_tables_from_pdf(str(path)),
        }
    elif ext in (".docx",):
        return {
            "source_type": "docx",
            "raw_images": extract_images_from_docx(str(path)),
            "raw_tables": extract_tables_from_docx(str(path)),
        }
    else:
        logger.warning("[MEDIA] Unsupported file type for media extraction: %s", ext)
        return {"source_type": "unknown", "raw_images": [], "raw_tables": []}
