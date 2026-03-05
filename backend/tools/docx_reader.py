from collections import Counter
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.logger import get_logger
from tools.tool_errors import ExtractionError, FileProcessingError

logger = get_logger(__name__)

GARBLE_RATIO = 0.6


def extract_docx_text(filepath: str) -> str:
    """
    Extract all text from a DOCX file as a plain string.

    Output format matches extract_pdf_text() for pipeline consistency.

    Args:
        filepath: Absolute path to the DOCX file.

    Returns:
        Concatenated text content from all paragraphs and tables.

    Raises:
        FileProcessingError: If the file does not exist or cannot be opened.
        ExtractionError: If the document contains no extractable text.
    """
    path = Path(filepath).resolve()
    if not path.exists():
        raise FileProcessingError(f"DOCX file not found: {path}")

    if path.suffix.lower() == ".doc":
        raise ExtractionError("Please convert .doc to .docx format first.")

    try:
        doc = Document(str(path))
    except Exception as e:
        msg = str(e).lower()
        if "password" in msg or "encrypted" in msg:
            raise ExtractionError("DOCX is password protected.")
        raise FileProcessingError(
            f"Could not open DOCX file. File may be corrupted: {e}"
        ) from e

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        parts.append(_extract_table_text(table))

    full_text = "\n".join(parts).strip()

    if not full_text:
        raise ExtractionError("Document contains no extractable text.")

    if _is_text_garbled(full_text):
        raise ExtractionError(
            "Extracted text appears corrupted or unreadable. "
            "The DOCX may contain encoding issues."
        )

    logger.info(
        "DOCX parsed | file=%s | paragraphs=%d | words=%d",
        path.name,
        len([p for p in doc.paragraphs if p.text.strip()]),
        len(full_text.split()),
    )
    return full_text


def extract_docx_structured(filepath: str) -> dict[str, Any]:
    """
    Extract rich structured data from a DOCX file.

    Args:
        filepath: Absolute path to the DOCX file.

    Returns:
        dict with keys:
            raw_text (str),
            paragraphs (list of paragraph metadata dicts),
            tables (list of table dicts),
            metadata (summary statistics dict).

    Raises:
        FileProcessingError: If the file does not exist or cannot be opened.
        ExtractionError: If the document contains no extractable text.
    """
    path = Path(filepath).resolve()
    if not path.exists():
        raise FileProcessingError(f"DOCX file not found: {path}")

    if path.suffix.lower() == ".doc":
        raise ExtractionError("Please convert .doc to .docx format first.")

    try:
        doc = Document(str(path))
    except Exception as e:
        msg = str(e).lower()
        if "password" in msg or "encrypted" in msg:
            raise ExtractionError("DOCX is password protected.")
        raise FileProcessingError(
            f"Could not open DOCX file. File may be corrupted: {e}"
        ) from e

    # Determine body font size from most common font size across all runs
    body_font_size = _detect_body_font_size(doc)

    paragraphs: list[dict[str, Any]] = []
    raw_parts: list[str] = []
    detected_styles: set[str] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        raw_parts.append(text)

        style_name = para.style.name if para.style else "Normal"
        detected_styles.add(style_name)

        is_heading, heading_level = _is_heading(para, body_font_size)

        bold = False
        italic = False
        font_name: Optional[str] = None
        font_size: Optional[float] = None
        for run in para.runs:
            if run.text.strip():
                bold = bool(run.bold)
                italic = bool(run.italic)
                if run.font.name:
                    font_name = run.font.name
                if run.font.size:
                    font_size = run.font.size.pt
                break

        paragraphs.append({
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "heading_level": heading_level,
            "bold": bold,
            "italic": italic,
            "font_name": font_name,
            "font_size": font_size,
            "alignment": _get_alignment_name(para.alignment),
        })

    tables_data: list[dict[str, Any]] = []
    for idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0
        content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables_data.append({"index": idx, "rows": rows, "cols": cols, "content": content})
        raw_parts.append(_extract_table_text(table))

    raw_text = "\n".join(raw_parts).strip()
    if not raw_text:
        raise ExtractionError("Document contains no extractable text.")

    total_words = sum(len(p["text"].split()) for p in paragraphs)

    has_images = any(
        run._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        )
        for para in doc.paragraphs
        for run in para.runs
    )

    metadata: dict[str, Any] = {
        "total_paragraphs": len(paragraphs),
        "total_words": total_words,
        "has_headings": any(p["is_heading"] for p in paragraphs),
        "has_tables": len(tables_data) > 0,
        "has_images": has_images,
        "detected_styles": sorted(detected_styles),
        "body_font_size": body_font_size,
    }

    logger.info(
        "DOCX structured | file=%s | paragraphs=%d | words=%d | tables=%d",
        path.name, len(paragraphs), total_words, len(tables_data),
    )

    return {
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "tables": tables_data,
        "metadata": metadata,
    }


def _detect_body_font_size(doc: Document) -> float:
    """
    Determine the most common font size used in the document body.
    Returns the mode font size across all non-empty runs, or 12.0 as default.
    """
    sizes: Counter[float] = Counter()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip() and run.font.size:
                sizes[run.font.size.pt] += 1
    return sizes.most_common(1)[0][0] if sizes else 12.0


def _extract_table_text(table) -> str:
    """Convert a python-docx Table to a pipe-delimited markdown-style text."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


def _get_alignment_name(alignment: Any) -> str:
    """Convert python-docx WD_ALIGN_PARAGRAPH enum to a human-readable string."""
    if alignment is None:
        return "left"
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(alignment, "left")


def _is_heading(paragraph, body_font_size: float = 12.0) -> tuple[bool, Optional[int]]:
    """
    Detect whether a paragraph is a heading and return its level.

    Detection rules (in priority order):
      1. Style name starts with "Heading" → extract level number
      2. Style name is "Title" or "Subtitle" → H1
      3. Font size > body_font_size → likely heading (level inferred from size)
      4. Bold + short (<60 chars) + no sentence-ending period → likely H2
      5. ALL CAPS short line (<60 chars) → likely H1

    Returns:
        (is_heading: bool, level: int | None)
    """
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()

    # Rule 1: "Heading 1", "Heading 2", etc.
    if style_name.lower().startswith("heading"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[-1].isdigit():
            return True, int(parts[-1])
        return True, 1

    # Rule 2: "Title" / "Subtitle" style
    if style_name.lower() in ("title", "subtitle"):
        return True, 1

    if not text:
        return False, None

    # Rule 3: Font size larger than body text
    run_sizes = [run.font.size.pt for run in paragraph.runs if run.text.strip() and run.font.size]
    if run_sizes:
        para_size = max(run_sizes)
        if para_size > body_font_size + 1:
            # Infer level from size relative to body
            delta = para_size - body_font_size
            level = 1 if delta >= 4 else (2 if delta >= 2 else 3)
            return True, level

    # Rule 4: Bold + short + no trailing period
    is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
    if is_bold and len(text) < 60 and not text.endswith("."):
        return True, 2

    # Rule 5: ALL CAPS short line
    if text.isupper() and len(text) < 60:
        return True, 1

    return False, None


def _is_text_garbled(text: str) -> bool:
    """
    Return True if the extracted text appears corrupted or unreadable.

    Detects excessive non-printable characters and near-zero word density.
    """
    if not text:
        return True
    printable = sum(1 for c in text if c.isprintable() or c in ("\n", "\t", "\r"))
    if printable / len(text) < GARBLE_RATIO:
        return True
    import re
    words = re.findall(r"[A-Za-z]{3,}", text)
    if len(text) > 200 and len(words) < len(text) / 50:
        return True
    return False


def get_docx_metadata(filepath: str) -> dict[str, Any]:
    """
    Return basic metadata for a DOCX file without full extraction.

    Raises:
        FileProcessingError: If file does not exist or cannot be opened.
    """
    path = Path(filepath).resolve()
    if not path.exists():
        raise FileProcessingError(f"DOCX file not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise FileProcessingError(f"Cannot open DOCX for metadata: {e}") from e

    styles = {p.style.name for p in doc.paragraphs if p.style}
    has_headings = any(s.lower().startswith("heading") for s in styles)

    return {
        "total_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
        "has_headings": has_headings,
        "has_tables": len(doc.tables) > 0,
        "detected_styles": sorted(styles),
    }


if __name__ == "__main__":
    test_path = "tests/sample.docx"
    try:
        text = extract_docx_text(test_path)
        assert len(text) > 100, "Extracted text too short"
        print(f"OK docx_reader (text): extracted {len(text)} characters")

        structured = extract_docx_structured(test_path)
        print(f"OK docx_reader (structured): {structured['metadata']}")
    except FileProcessingError:
        print(f"SKIP no sample at {test_path}")
    except ExtractionError as e:
        print(f"ExtractionError: {e}")

    try:
        extract_docx_text("nonexistent_file_xyz.docx")
    except FileProcessingError:
        print("OK docx_reader: FileProcessingError handled correctly")
